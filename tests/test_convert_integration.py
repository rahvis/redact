"""
Integration tests for the Convert menu group and OCR handler cores
(workonward_read.handlers.convert): headless request execution for the four
PDF conversions, images-to-PDF, the redacted-raster OCR flow and the
tesseract settings JSON.

Licensed under GPL-3.0
(c) 2024 - 2026 Björn Seipel
Acrobat-suite additions (c) 2026 CoverUP contributors
"""

import json
import os

import pytest
import pypdfium2 as pdfium
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader

import fixtures
from workonward_read import ocr
from workonward_read.handlers import convert as convert_handlers
from workonward_read.image_container import ImageContainer
from workonward_read.state import AppState


TESSERACT = ocr.find_tesseract()
needs_tesseract = pytest.mark.skipif(
    TESSERACT is None, reason='tesseract binary not installed')

DPI = 200


def _state(file_path=None, password=None, images=None):
    state = AppState()
    state.file_path = file_path
    state.source_password = password
    if images is not None:
        state.images = images
    return state


def _extract_text_pdfium(path):
    doc = pdfium.PdfDocument(str(path))
    try:
        chunks = []
        for page in doc:
            textpage = page.get_textpage()
            chunks.append(textpage.get_text_range())
            textpage.close()
            page.close()
        return '\n'.join(chunks)
    finally:
        doc.close()


# ---------------------------------------------------------------------------
# HANDLERS dict shape
# ---------------------------------------------------------------------------

def test_handlers_dict_shape():
    from workonward_read.dialogs.common import not_yet
    expected = {
        'MENU_CONVERT_IMAGES', 'MENU_CONVERT_TEXT', 'MENU_CONVERT_WORD',
        'MENU_CONVERT_HTML', 'MENU_IMAGES_TO_PDF', 'MENU_OCR',
    }
    assert set(convert_handlers.HANDLERS) == expected
    for key, handler in convert_handlers.HANDLERS.items():
        assert callable(handler), key
        assert handler is not not_yet, f'{key} still maps to the stub'


# ---------------------------------------------------------------------------
# run_convert: the four PDF conversion targets
# ---------------------------------------------------------------------------

def test_convert_to_images(tmp_path):
    pdf = fixtures.make_pdf(tmp_path / 'doc.pdf', pages=2)
    out_dir = tmp_path / 'imgs'
    request = {'target': 'images', 'use_loaded': False, 'input_path': pdf,
               'output': str(out_dir), 'fmt': 'PNG', 'dpi': 100}
    progress = []
    paths = convert_handlers.run_convert(
        request, _state(), progress_cb=lambda pct, msg: progress.append(pct))

    assert len(paths) == 2
    for path in paths:
        assert os.path.isfile(path)
        with Image.open(path) as img:
            assert img.format == 'PNG'
    assert progress and progress[-1] == 100


def test_convert_to_images_jpeg(tmp_path):
    pdf = fixtures.make_pdf(tmp_path / 'doc.pdf', pages=1)
    request = {'target': 'images', 'use_loaded': False, 'input_path': pdf,
               'output': str(tmp_path / 'jpgs'), 'fmt': 'JPEG', 'dpi': 72}
    paths = convert_handlers.run_convert(request, _state())
    assert len(paths) == 1 and paths[0].endswith('.jpg')
    with Image.open(paths[0]) as img:
        assert img.format == 'JPEG'


def test_convert_to_text(tmp_path):
    text = 'Alpha bravo charlie delta'
    pdf = fixtures.make_text_pdf(tmp_path / 'doc.pdf', text=text)
    out = str(tmp_path / 'doc.txt')
    request = {'target': 'text', 'use_loaded': False, 'input_path': pdf,
               'output': out}
    paths = convert_handlers.run_convert(request, _state())

    assert paths == [out]
    with open(out, encoding='utf-8') as fh:
        assert text in fh.read()


def test_convert_to_word(tmp_path):
    from docx import Document

    text = 'Word conversion paragraph content'
    pdf = fixtures.make_text_pdf(tmp_path / 'doc.pdf', text=text)
    out = str(tmp_path / 'doc.docx')
    request = {'target': 'word', 'use_loaded': False, 'input_path': pdf,
               'output': out}
    paths = convert_handlers.run_convert(request, _state())

    assert paths == [out] and os.path.isfile(out)
    document = Document(out)
    combined = '\n'.join(p.text for p in document.paragraphs)
    assert text in combined


def test_convert_to_html(tmp_path):
    text = 'HTML conversion body text'
    pdf = fixtures.make_text_pdf(tmp_path / 'doc.pdf', text=text)
    out = str(tmp_path / 'doc.html')
    request = {'target': 'html', 'use_loaded': False, 'input_path': pdf,
               'output': out, 'embed_page_images': False}
    paths = convert_handlers.run_convert(request, _state())

    assert paths == [out]
    with open(out, encoding='utf-8') as fh:
        content = fh.read()
    assert text in content
    assert 'data:image/jpeg;base64,' not in content


def test_convert_to_html_with_embedded_page_images(tmp_path):
    pdf = fixtures.make_text_pdf(tmp_path / 'doc.pdf', text='embedded run')
    out = str(tmp_path / 'doc.html')
    request = {'target': 'html', 'use_loaded': False, 'input_path': pdf,
               'output': out, 'embed_page_images': True}
    convert_handlers.run_convert(request, _state())
    with open(out, encoding='utf-8') as fh:
        assert 'data:image/jpeg;base64,' in fh.read()


def test_convert_loaded_file_uses_source_password(tmp_path):
    pdf = fixtures.make_encrypted_pdf(tmp_path / 'enc.pdf',
                                      user_password='secret', pages=1)
    out = str(tmp_path / 'enc.txt')
    request = {'target': 'text', 'use_loaded': True, 'input_path': None,
               'output': out}
    state = _state(file_path=pdf, password='secret')
    paths = convert_handlers.run_convert(request, state)

    assert paths == [out]
    with open(out, encoding='utf-8') as fh:
        assert 'Encrypted page 1' in fh.read()


def test_convert_rejects_missing_input():
    with pytest.raises(ValueError):
        convert_handlers.run_convert(
            {'target': 'text', 'use_loaded': False, 'input_path': '',
             'output': 'x.txt'}, _state())
    with pytest.raises(ValueError):
        convert_handlers.run_convert(
            {'target': 'text', 'use_loaded': True, 'input_path': None,
             'output': 'x.txt'}, _state(file_path=None))


def test_convert_rejects_unknown_target(tmp_path):
    pdf = fixtures.make_text_pdf(tmp_path / 'doc.pdf')
    with pytest.raises(ValueError):
        convert_handlers.run_convert(
            {'target': 'bogus', 'use_loaded': False, 'input_path': pdf,
             'output': str(tmp_path / 'x')}, _state())


# ---------------------------------------------------------------------------
# run_images_to_pdf
# ---------------------------------------------------------------------------

def test_images_to_pdf(tmp_path):
    img_a = fixtures.make_image(tmp_path / 'a.png', size=(400, 300),
                                text='first')
    img_b = fixtures.make_image(tmp_path / 'b.jpg', size=(300, 400),
                                text='second')
    out = str(tmp_path / 'combined.pdf')
    result = convert_handlers.run_images_to_pdf(
        {'image_paths': [img_a, img_b], 'output': out})

    assert result == out
    reader = PdfReader(out)
    assert len(reader.pages) == 2


def test_images_to_pdf_empty_list_raises(tmp_path):
    with pytest.raises(ValueError):
        convert_handlers.run_images_to_pdf(
            {'image_paths': [], 'output': str(tmp_path / 'x.pdf')})


# ---------------------------------------------------------------------------
# Settings JSON (custom tesseract path)
# ---------------------------------------------------------------------------

def test_settings_load_missing_file_returns_empty(tmp_path):
    assert convert_handlers.load_settings(str(tmp_path)) == {}
    assert convert_handlers.get_saved_tesseract_path(str(tmp_path)) is None


def test_settings_tesseract_path_roundtrip(tmp_path):
    datadir = str(tmp_path)
    path = convert_handlers.save_tesseract_path('/opt/custom/tesseract',
                                                datadir)
    assert path == convert_handlers.settings_path(datadir)
    assert os.path.isfile(path)

    # Well-formed JSON with the expected key.
    with open(path, encoding='utf-8') as fh:
        data = json.load(fh)
    assert data[convert_handlers.TESSERACT_PATH_KEY] == '/opt/custom/tesseract'
    assert (convert_handlers.get_saved_tesseract_path(datadir)
            == '/opt/custom/tesseract')


def test_settings_save_preserves_other_keys(tmp_path):
    datadir = str(tmp_path)
    convert_handlers.save_settings({'other': 1}, datadir)
    convert_handlers.save_tesseract_path('/usr/bin/tesseract', datadir)
    settings = convert_handlers.load_settings(datadir)
    assert settings['other'] == 1
    assert settings[convert_handlers.TESSERACT_PATH_KEY] == '/usr/bin/tesseract'


def test_settings_corrupt_file_returns_empty(tmp_path):
    datadir = str(tmp_path)
    with open(convert_handlers.settings_path(datadir), 'w',
              encoding='utf-8') as fh:
        fh.write('{not json')
    assert convert_handlers.load_settings(datadir) == {}


# ---------------------------------------------------------------------------
# OCR: current-document (redacted raster) flow
# ---------------------------------------------------------------------------

VISIBLE_TEXT = 'HELLO WORLD'
HIDDEN_TEXT = 'TOPSECRET'


def _redacted_container(visible=VISIBLE_TEXT, hidden=HIDDEN_TEXT, dpi=DPI):
    """
    A4-sized page container with large visible text and large hidden text
    covered by a black 'redact' annotation.
    """
    w_px, h_px = int(8.27 * dpi), int(11.69 * dpi)
    img = Image.new('RGB', (w_px, h_px), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default(size=80)

    draw.text((w_px // 8, h_px // 4), visible, fill=(0, 0, 0), font=font)
    hidden_pos = (w_px // 8, h_px // 2)
    draw.text(hidden_pos, hidden, fill=(0, 0, 0), font=font)
    hidden_bbox = draw.textbbox(hidden_pos, hidden, font=font)

    # ImageContainer.size is (width_pt, height_pt), like document_loader.
    container = ImageContainer(img, (w_px * 72 / dpi, h_px * 72 / dpi))
    container.add_annotation('redact', {
        'p1': [hidden_bbox[0] - 20, hidden_bbox[1] - 20],
        'p2': [hidden_bbox[2] + 20, hidden_bbox[3] + 20],
        'fill': 'black',
    })
    return container, hidden_bbox


def test_iter_redacted_pages_burns_annotations_headlessly():
    """The finalize path must blacken the redacted area (no tesseract needed)."""
    container, hidden_bbox = _redacted_container()
    try:
        pages = list(convert_handlers.iter_redacted_pages([container]))
        assert len(pages) == 1
        # The yielded image was closed after iteration; re-render to inspect.
        burned = container.finalized_image('PIL')
        center = ((hidden_bbox[0] + hidden_bbox[2]) // 2,
                  (hidden_bbox[1] + hidden_bbox[3]) // 2)
        assert burned.getpixel(center) == (0, 0, 0)
        # Original page image is untouched (text pixels still present).
        assert container.image.getpixel(center) != (255, 255, 255)
        burned.close()
    finally:
        container.close()


@needs_tesseract
def test_ocr_current_doc_hides_redacted_text(tmp_path):
    """OCR of the loaded document must not contain covered text but must
    contain the visible text."""
    container, _bbox = _redacted_container()
    state = _state(file_path='/nonexistent/loaded.pdf', images=[container])
    out = str(tmp_path / 'redacted_ocr.pdf')

    progress = []
    result = convert_handlers.run_ocr(
        {'use_loaded': True, 'input_path': None, 'lang': 'eng',
         'output': out},
        state, tess_path=TESSERACT,
        progress_cb=lambda pct, msg: progress.append(pct))

    assert result == out and os.path.isfile(out)
    text = _extract_text_pdfium(out).upper()
    assert 'HELLO' in text
    assert 'WORLD' in text
    assert 'TOPSECRET' not in text
    assert 'SECRET' not in text
    assert progress and progress[-1] == 100

    # Page geometry preserved (A4 within a point).
    reader = PdfReader(out)
    box = reader.pages[0].mediabox
    assert abs(float(box.width) - 8.27 * 72) < 1.5
    assert abs(float(box.height) - 11.69 * 72) < 1.5
    container.close()


@needs_tesseract
def test_ocr_picked_file(tmp_path):
    """OCR of a picked scanned PDF goes through ocr.make_searchable_pdf."""
    scan = str(tmp_path / 'scan.pdf')
    img = Image.new('RGB', (int(8.27 * DPI), int(11.69 * DPI)),
                    (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default(size=80)
    draw.text((img.width // 8, img.height // 3), VISIBLE_TEXT,
              fill=(0, 0, 0), font=font)
    import io
    from fpdf import FPDF
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=90)
    buf.seek(0)
    pdf = FPDF(unit='pt', format=(img.width * 72 / DPI,
                                  img.height * 72 / DPI))
    pdf.add_page()
    pdf.image(buf, x=0, y=0, w=pdf.w)
    pdf.output(scan)
    img.close()

    out = str(tmp_path / 'scan_ocr.pdf')
    result = convert_handlers.run_ocr(
        {'use_loaded': False, 'input_path': scan, 'lang': 'eng',
         'output': out},
        _state(), tess_path=TESSERACT)

    assert result == out
    text = _extract_text_pdfium(out).upper()
    assert 'HELLO' in text and 'WORLD' in text


def test_run_ocr_rejects_empty_state_and_input(tmp_path):
    with pytest.raises(ValueError):
        convert_handlers.run_ocr(
            {'use_loaded': True, 'output': str(tmp_path / 'x.pdf')},
            _state(images=[]))
    with pytest.raises(ValueError):
        convert_handlers.run_ocr(
            {'use_loaded': False, 'input_path': '',
             'output': str(tmp_path / 'x.pdf')},
            _state())
