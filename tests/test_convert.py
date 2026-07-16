"""
Tests for workonward_read.convert — export/conversion tools.

All fixtures are synthesized at test time (see tests/fixtures.py).

Licensed under GPL-3.0. (c) 2024 - 2026 Björn Seipel
Acrobat-suite additions (c) 2026 CoverUP contributors
"""

import csv
import os

import pytest
from fpdf import FPDF
from PIL import Image
from pypdf import PdfReader

import fixtures
from workonward_read import convert


A4_PT = (595.28, 841.89)


def _noise_image(size=(800, 1000)):
    """Random-noise RGB image: compresses badly, ideal for size assertions."""
    width, height = size
    return Image.frombytes("RGB", (width, height), os.urandom(width * height * 3))


def _make_image_heavy_pdf(path, text="RASTER SOURCE TEXT"):
    """PDF with a text layer plus a large losslessly-embedded noise image."""
    noise = _noise_image()
    png_path = str(path) + ".noise.png"
    noise.save(png_path, format="PNG")
    noise.close()

    pdf = FPDF(unit="pt", format=A4_PT)
    pdf.set_auto_page_break(False)
    pdf.add_page()
    pdf.set_font("helvetica", size=16)
    pdf.set_xy(72, 40)
    pdf.cell(400, 20, text)
    pdf.image(png_path, x=20, y=80, w=550, h=680)
    pdf.output(str(path))
    os.remove(png_path)
    return str(path)


def _extracted_text(pdf_path, tmp_path, name="extract.txt", password=None):
    out = convert.pdf_to_text(pdf_path, str(tmp_path / name), password=password)
    with open(out, encoding="utf-8") as fh:
        return fh.read()


# ---------------------------------------------------------------------------
# pdf_to_images
# ---------------------------------------------------------------------------

class TestPdfToImages:
    def test_roundtrip_page_count_and_dims(self, tmp_path):
        pdf_path = fixtures.make_pdf(tmp_path / "doc.pdf", pages=3, size=A4_PT)
        out_dir = tmp_path / "imgs"
        dpi = 100

        progress = []
        written = convert.pdf_to_images(
            pdf_path, str(out_dir), dpi=dpi,
            progress_cb=lambda pct, msg: progress.append((pct, msg)),
        )

        assert len(written) == 3
        for n, path in enumerate(written, start=1):
            assert os.path.isfile(path)
            assert os.path.basename(path) == f"doc_page_{n}.png"
        with Image.open(written[0]) as img:
            assert abs(img.width - A4_PT[0] * dpi / 72) <= 2
            assert abs(img.height - A4_PT[1] * dpi / 72) <= 2
        assert progress and progress[-1][0] == 100

    def test_page_subset_and_jpeg(self, tmp_path):
        pdf_path = fixtures.make_pdf(tmp_path / "doc.pdf", pages=3)
        written = convert.pdf_to_images(
            pdf_path, str(tmp_path / "imgs"), fmt="JPEG", dpi=72, pages=[1]
        )
        assert len(written) == 1
        assert os.path.basename(written[0]) == "doc_page_2.jpg"
        with Image.open(written[0]) as img:
            assert img.format == "JPEG"

    def test_invalid_page_index(self, tmp_path):
        pdf_path = fixtures.make_pdf(tmp_path / "doc.pdf", pages=2)
        with pytest.raises(ValueError):
            convert.pdf_to_images(pdf_path, str(tmp_path / "imgs"), pages=[5])

    def test_missing_file(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            convert.pdf_to_images(str(tmp_path / "nope.pdf"), str(tmp_path))


# ---------------------------------------------------------------------------
# pdf_to_text
# ---------------------------------------------------------------------------

class TestPdfToText:
    def test_contains_text_and_form_feeds(self, tmp_path):
        text = "The quick brown fox visits Zürich."
        pdf_path = fixtures.make_text_pdf(tmp_path / "doc.pdf", text=text, pages=2)
        out = convert.pdf_to_text(pdf_path, str(tmp_path / "doc.txt"))

        assert out == str(tmp_path / "doc.txt")
        content = open(out, encoding="utf-8").read()
        assert text in content
        assert len(content.split("\f")) == 2

    def test_returns_path(self, tmp_path):
        pdf_path = fixtures.make_pdf(tmp_path / "doc.pdf", pages=1)
        out = convert.pdf_to_text(pdf_path, str(tmp_path / "out.txt"))
        assert os.path.isfile(out)


# ---------------------------------------------------------------------------
# pdf_to_docx
# ---------------------------------------------------------------------------

class TestPdfToDocx:
    def test_paragraphs_and_text(self, tmp_path):
        from docx import Document

        long_text = (
            "First paragraph line one that is long enough to wrap nicely.\n\n"
            "Second paragraph after a blank gap in the source text."
        )
        pdf_path = fixtures.make_pdf(
            tmp_path / "doc.pdf", pages=2, texts=[long_text, "Second page text."]
        )
        out = convert.pdf_to_docx(pdf_path, str(tmp_path / "doc.docx"))

        document = Document(out)
        texts = [p.text for p in document.paragraphs if p.text.strip()]
        assert len(texts) > 0
        joined = " ".join(texts)
        assert "First paragraph line one" in joined
        assert "Second page text." in joined


# ---------------------------------------------------------------------------
# pdf_to_html
# ---------------------------------------------------------------------------

class TestPdfToHtml:
    def test_sections_and_escaping(self, tmp_path):
        tricky = "Tom & Jerry <b>not bold</b>"
        pdf_path = fixtures.make_pdf(
            tmp_path / "doc.pdf", pages=2, texts=[tricky, "Plain page two"]
        )
        out = convert.pdf_to_html(pdf_path, str(tmp_path / "doc.html"))

        content = open(out, encoding="utf-8").read()
        assert content.count("<section") == 2
        assert "Tom &amp; Jerry" in content
        assert "&lt;b&gt;not bold&lt;/b&gt;" in content
        assert "<b>not bold</b>" not in content
        assert "Plain page two" in content

    def test_embedded_page_images(self, tmp_path):
        pdf_path = fixtures.make_pdf(tmp_path / "doc.pdf", pages=1)
        out = convert.pdf_to_html(
            pdf_path, str(tmp_path / "doc.html"), embed_page_images=True
        )
        content = open(out, encoding="utf-8").read()
        assert "data:image/jpeg;base64," in content


# ---------------------------------------------------------------------------
# pdf_to_csv_text
# ---------------------------------------------------------------------------

class TestPdfToCsvText:
    def test_rows_written(self, tmp_path):
        pdf_path = fixtures.make_pdf(
            tmp_path / "doc.pdf", pages=1, texts=["Alpha  Beta  Gamma"]
        )
        out = convert.pdf_to_csv_text(pdf_path, str(tmp_path / "doc.csv"))

        assert os.path.isfile(out)
        with open(out, encoding="utf-8", newline="") as fh:
            rows = list(csv.reader(fh))
        assert rows, "expected at least one CSV row"
        flattened = [cell for row in rows for cell in row]
        assert any("Alpha" in cell for cell in flattened)


# ---------------------------------------------------------------------------
# images_to_pdf
# ---------------------------------------------------------------------------

class TestImagesToPdf:
    def test_page_sizes_and_count(self, tmp_path):
        small = fixtures.make_image(tmp_path / "small.png", size=(400, 300))
        big = fixtures.make_image(tmp_path / "big.jpg", size=(3000, 2000))

        out = convert.images_to_pdf([small, big], str(tmp_path / "out.pdf"))

        reader = PdfReader(out)
        assert len(reader.pages) == 2

        # Small image keeps its size: int(px / 200 * 72)
        page0 = reader.pages[0]
        assert float(page0.mediabox.width) == pytest.approx(int(400 / 200 * 72))
        assert float(page0.mediabox.height) == pytest.approx(int(300 / 200 * 72))

        # Large landscape image is scaled to fit A4 at 200 PPI
        a4_long_pt = round(11.693 * 200) / 200 * 72  # ~841.9 pt
        page1 = reader.pages[1]
        assert float(page1.mediabox.width) <= a4_long_pt + 1
        assert float(page1.mediabox.width) > float(page1.mediabox.height)

    def test_empty_list_raises(self, tmp_path):
        with pytest.raises(ValueError):
            convert.images_to_pdf([], str(tmp_path / "out.pdf"))


# ---------------------------------------------------------------------------
# compress_pdf_raster
# ---------------------------------------------------------------------------

class TestCompressRaster:
    def test_shrinks_and_flattens_text(self, tmp_path):
        pdf_path = _make_image_heavy_pdf(tmp_path / "heavy.pdf")
        out = str(tmp_path / "small.pdf")

        progress = []
        convert.compress_pdf_raster(
            pdf_path, out, dpi=72, jpeg_quality=60,
            progress_cb=lambda pct, msg: progress.append(pct),
        )

        assert os.path.getsize(out) < os.path.getsize(pdf_path)
        assert progress and progress[-1] == 100

        # Original text layer is gone (flattened to an image)
        remaining = _extracted_text(out, tmp_path).replace("\f", "").strip()
        assert remaining == ""

        # Page size in points is preserved
        reader = PdfReader(out)
        assert float(reader.pages[0].mediabox.width) == pytest.approx(A4_PT[0], abs=1)
        assert float(reader.pages[0].mediabox.height) == pytest.approx(A4_PT[1], abs=1)


# ---------------------------------------------------------------------------
# compress_pdf_lossless
# ---------------------------------------------------------------------------

class TestCompressLossless:
    def test_shrinks_and_keeps_text(self, tmp_path):
        text = "RASTER SOURCE TEXT"
        pdf_path = _make_image_heavy_pdf(tmp_path / "heavy.pdf", text=text)
        out = str(tmp_path / "small.pdf")

        result = convert.compress_pdf_lossless(pdf_path, out, image_quality=40)

        assert result["before_bytes"] == os.path.getsize(pdf_path)
        assert result["after_bytes"] == os.path.getsize(out)
        assert result["before_bytes"] > result["after_bytes"]

        # Text layer survives
        assert text in _extracted_text(out, tmp_path)

    def test_plain_text_pdf_survives(self, tmp_path):
        pdf_path = fixtures.make_text_pdf(tmp_path / "plain.pdf", pages=2)
        out = str(tmp_path / "out.pdf")
        result = convert.compress_pdf_lossless(pdf_path, out)
        assert os.path.isfile(out)
        assert set(result) == {"before_bytes", "after_bytes"}


# ---------------------------------------------------------------------------
# batch_apply
# ---------------------------------------------------------------------------

class TestBatchApply:
    def test_happy_and_error_paths(self, tmp_path):
        good1 = fixtures.make_text_pdf(tmp_path / "one.pdf", text="Batch one")
        good2 = fixtures.make_text_pdf(tmp_path / "two.pdf", text="Batch two")
        bad = tmp_path / "broken.pdf"
        bad.write_bytes(b"this is not a pdf")

        out_dir = tmp_path / "batch_out"
        progress = []
        results = convert.batch_apply(
            convert.pdf_to_text, [good1, str(bad), good2], str(out_dir),
            progress_cb=lambda pct, msg: progress.append((pct, msg)),
            out_ext=".txt",
        )

        assert len(results) == 3
        assert [r[0] for r in results] == [good1, str(bad), good2]

        in1, out1, err1 = results[0]
        assert err1 is None and os.path.isfile(out1) and out1.endswith("one.txt")
        assert "Batch one" in open(out1, encoding="utf-8").read()

        in_bad, out_bad, err_bad = results[1]
        assert out_bad is None
        assert isinstance(err_bad, str) and err_bad

        in2, out2, err2 = results[2]
        assert err2 is None and os.path.isfile(out2)

        assert progress and progress[-1][0] == 100

    def test_stem_collision_gets_suffix(self, tmp_path):
        dir_a = tmp_path / "a"
        dir_b = tmp_path / "b"
        dir_a.mkdir()
        dir_b.mkdir()
        pdf_a = fixtures.make_text_pdf(dir_a / "same.pdf", text="From A")
        pdf_b = fixtures.make_text_pdf(dir_b / "same.pdf", text="From B")

        results = convert.batch_apply(
            convert.pdf_to_text, [pdf_a, pdf_b], str(tmp_path / "out"),
            out_ext=".txt",
        )
        outputs = [r[1] for r in results]
        assert None not in outputs
        assert len(set(outputs)) == 2
        assert all(os.path.isfile(p) for p in outputs)


# ---------------------------------------------------------------------------
# Encrypted inputs
# ---------------------------------------------------------------------------

class TestEncryptedInputs:
    def test_pdf_to_text_with_password(self, tmp_path):
        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret", pages=2)
        content = _extracted_text(pdf_path, tmp_path, password="secret")
        assert "Encrypted page 1" in content

    def test_pdf_to_text_without_password_raises(self, tmp_path):
        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret")
        with pytest.raises(ValueError, match="password"):
            convert.pdf_to_text(pdf_path, str(tmp_path / "out.txt"))

    def test_pdf_to_images_with_password(self, tmp_path):
        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret", pages=1)
        written = convert.pdf_to_images(
            pdf_path, str(tmp_path / "imgs"), dpi=72, password="secret"
        )
        assert len(written) == 1 and os.path.isfile(written[0])

    def test_pdf_to_docx_with_password(self, tmp_path):
        from docx import Document

        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret", pages=1)
        out = convert.pdf_to_docx(
            pdf_path, str(tmp_path / "enc.docx"), password="secret"
        )
        document = Document(out)
        joined = " ".join(p.text for p in document.paragraphs)
        assert "Encrypted page 1" in joined

    def test_pdf_to_docx_wrong_password_raises(self, tmp_path):
        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret")
        with pytest.raises(ValueError, match="password"):
            convert.pdf_to_docx(
                pdf_path, str(tmp_path / "enc.docx"), password="wrong"
            )

    def test_compress_lossless_with_password(self, tmp_path):
        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret", pages=1)
        out = str(tmp_path / "out.pdf")
        result = convert.compress_pdf_lossless(pdf_path, out, password="secret")
        assert result["after_bytes"] > 0

    def test_compress_lossless_missing_password_raises(self, tmp_path):
        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret")
        with pytest.raises(ValueError, match="password"):
            convert.compress_pdf_lossless(pdf_path, str(tmp_path / "out.pdf"))

    def test_compress_raster_with_password(self, tmp_path):
        pdf_path = fixtures.make_encrypted_pdf(tmp_path / "enc.pdf", "secret", pages=1)
        out = convert.compress_pdf_raster(
            pdf_path, str(tmp_path / "out.pdf"), dpi=72, password="secret"
        )
        assert os.path.getsize(out) > 0
